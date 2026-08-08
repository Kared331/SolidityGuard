"""Report generation with fixed fuzz parsing (2.4) and LLM data separation (4.17)."""

import json
import logging
import os
from datetime import datetime, timezone

from jinja2 import Environment, FileSystemLoader
from sqlalchemy.orm import Session

from app.models import Detection, FalsePositiveFeedback, FuzzingResult, LLMAuditResult
from app.services.infra.storage import get_report_dir
from app.llm.sync_wrapper import chat_completion

logger = logging.getLogger("solidguard.services.report_generator")


def aggregate_findings(project_id: int, session: Session) -> dict:
    """Aggregate findings from Slither, Fuzzing, and LLM audit."""
    false_positive_refs = {
        row[0]
        for row in session.query(FalsePositiveFeedback.detection_ref)
        .filter(FalsePositiveFeedback.project_id == project_id)
        .all()
    }

    from app.models import AnalysisResult

    detections = (
        session.query(Detection)
        .join(AnalysisResult, Detection.analysis_result_id == AnalysisResult.id)
        .filter(AnalysisResult.project_id == project_id)
        .all()
    )

    slither_findings = []
    for det in detections:
        if det.detection_ref in false_positive_refs:
            continue
        element = det.element_json or {}
        slither_findings.append({
            "check_name": det.check_name,
            "description": det.description,
            "impact": det.impact or "Unknown",
            "code_location": element.get("source_mapping", {}).get("filename_relative", "N/A"),
            "fix_suggestion": "",
            "gas_optimization": "",
        })

    # 2.4: Correctly handle fuzz failures_json which is a LIST, not a dict
    fuzzing_results = (
        session.query(FuzzingResult)
        .filter(FuzzingResult.project_id == project_id)
        .all()
    )
    fuzzing_findings = []
    for fr in fuzzing_results:
        failures = fr.failures_json
        if failures is None:
            continue
        # failures_json is a list of {"test_name": ..., "counterexample": ...}
        if isinstance(failures, list):
            for item in failures:
                fuzzing_findings.append({
                    "title": item.get("test_name", "unknown"),
                    "description": str(item.get("counterexample", "No counterexample")),
                    "severity": "High",
                    "code_location": item.get("test_name", "N/A"),
                    "fix_suggestion": "",
                    "gas_optimization": "",
                })
        elif isinstance(failures, dict):
            # Legacy dict format – iterate key-value pairs
            for name, detail in failures.items():
                fuzzing_findings.append({
                    "title": name,
                    "description": str(detail),
                    "severity": "High",
                    "code_location": name,
                    "fix_suggestion": "",
                    "gas_optimization": "",
                })

    llm_results = (
        session.query(LLMAuditResult)
        .filter(LLMAuditResult.project_id == project_id)
        .all()
    )
    llm_findings = []
    for lr in llm_results:
        llm_findings.append({
            "contract_name": lr.contract_name,
            "function_name": lr.function_name,
            "vulnerability_description": lr.vulnerability_description,
            "severity": lr.severity,
            "suggested_fix": lr.suggested_fix or "",
            "gas_optimization": lr.gas_optimization or "",
        })

    return {
        "slither_findings": slither_findings,
        "fuzzing_findings": fuzzing_findings,
        "llm_findings": llm_findings,
    }


_EXPECTED_KEYS = {"slither_findings", "fuzzing_findings", "llm_findings"}
_BATCH_SIZE = 5


def _polish_single_batch(batch_items: list, key: str) -> list:
    """Polish a single batch of findings via LLM."""
    system_prompt = (
        "You are a professional Solidity security audit report writer. "
        "Enhance the following findings by improving descriptions, adding context, "
        "and ensuring severity ratings are consistent. "
        "Return a JSON array of the polished findings in the same order."
    )
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": json.dumps(batch_items, indent=2)},
    ]
    response_text, _ = chat_completion(messages)
    polished = json.loads(response_text)
    if not isinstance(polished, list):
        raise ValueError(f"Expected list, got {type(polished).__name__}")
    if len(polished) != len(batch_items):
        logger.warning(
            "Polished batch size mismatch for %s: expected %d, got %d; using raw",
            key, len(batch_items), len(polished),
        )
        return batch_items
    return polished


def polish_with_llm(findings: dict) -> dict:
    """Polish findings with LLM in batches (max 5 per batch).

    Each category (slither, fuzzing, llm) is processed independently
    in batches to keep token usage bounded.  If any batch fails, the
    raw findings for that batch are preserved and other batches are
    unaffected.
    """
    polished = {}

    for key in _EXPECTED_KEYS:
        items = findings.get(key, [])
        if not items:
            polished[key] = []
            continue

        batched: list = []
        for i in range(0, len(items), _BATCH_SIZE):
            batch = items[i : i + _BATCH_SIZE]
            try:
                batch_polished = _polish_single_batch(batch, key)
                batched.extend(batch_polished)
            except (json.JSONDecodeError, ValueError, Exception):
                logger.warning(
                    "LLM polishing failed for %s batch %d-%d, using raw findings",
                    key, i, i + len(batch),
                )
                batched.extend(batch)

        polished[key] = batched

    # Validate final structure
    for key in _EXPECTED_KEYS:
        if not isinstance(polished.get(key), list):
            logger.warning(
                "LLM polishing result key '%s' is not a list, using raw findings", key
            )
            return findings

    return polished


def generate_html(project_id: int, title: str, findings: dict) -> str:
    template_dir = os.path.join(os.path.dirname(__file__), "templates")
    env = Environment(loader=FileSystemLoader(template_dir))
    template = env.get_template("report.html.jinja2")

    total = (
        len(findings.get("slither_findings", []))
        + len(findings.get("fuzzing_findings", []))
        + len(findings.get("llm_findings", []))
    )

    html = template.render(
        title=title,
        generated_date=datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC"),
        project_id=project_id,
        total_findings=total,
        slither_findings=findings.get("slither_findings", []),
        fuzzing_findings=findings.get("fuzzing_findings", []),
        llm_findings=findings.get("llm_findings", []),
    )

    report_dir = get_report_dir(project_id)
    os.makedirs(report_dir, exist_ok=True)
    html_path = os.path.join(report_dir, "report.html")
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html)

    return html_path


def generate_pdf(html_path: str) -> str:
    from weasyprint import HTML

    pdf_path = html_path.replace(".html", ".pdf")
    HTML(filename=html_path).write_pdf(pdf_path)
    return pdf_path


def generate_word(findings: dict, title: str, project_id: int) -> str:
    from docx import Document
    from docx.shared import Pt, Inches

    doc = Document()

    doc.add_heading(title, level=0)
    doc.add_paragraph(
        f"Generated on {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
    )

    total = (
        len(findings.get("slither_findings", []))
        + len(findings.get("fuzzing_findings", []))
        + len(findings.get("llm_findings", []))
    )
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(f"Total findings: {total}")

    sections = [
        ("Slither Static Analysis Findings", "slither_findings",
         ["check_name", "description", "impact", "code_location", "fix_suggestion", "gas_optimization"]),
        ("Fuzzing Findings", "fuzzing_findings",
         ["title", "description", "severity", "code_location", "fix_suggestion", "gas_optimization"]),
        ("LLM Audit Findings", "llm_findings",
         ["contract_name", "function_name", "vulnerability_description", "severity", "suggested_fix", "gas_optimization"]),
    ]

    for section_title, key, fields in sections:
        doc.add_heading(section_title, level=1)
        items = findings.get(key, [])
        if not items:
            doc.add_paragraph("No findings detected.")
            continue
        for i, item in enumerate(items, 1):
            doc.add_paragraph(f"Finding #{i}", style="Heading 2")
            for field in fields:
                value = item.get(field, "")
                if value:
                    p = doc.add_paragraph()
                    run = p.add_run(f"{field}: ")
                    run.bold = True
                    p.add_run(str(value))

    docx_path = os.path.join(get_report_dir(project_id), "report.docx")
    os.makedirs(os.path.dirname(docx_path), exist_ok=True)
    doc.save(docx_path)
    return docx_path
